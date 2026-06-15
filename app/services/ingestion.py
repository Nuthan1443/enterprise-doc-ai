import logging
import pandas as pd
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class IngestedDocument:
    """
    Output contract of the ingestion layer.
    Every downstream service receives this — never raw file bytes.
    """
    content: str
    metadata: dict = field(default_factory=dict)


class DocumentIngestionService:
    """
    Converts uploaded files into IngestedDocument objects.
    Supports: PDF, DOCX, TXT, CSV
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv"}

    def ingest(self, file_path: str | Path) -> IngestedDocument:
        """
        Main entry point. Routes to the correct parser by file extension.
        Raises ValueError for unsupported formats.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported format: '{extension}'. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        logger.info(f"Ingesting {extension} file: {path.name}")

        parsers = {
            ".pdf":  self._parse_pdf,
            ".docx": self._parse_docx,
            ".txt":  self._parse_txt,
            ".csv":  self._parse_csv,
        }

        content = parsers[extension](path)

        return IngestedDocument(
            content=content,
            metadata={
                "source":    path.name,
                "extension": extension,
                "file_size_bytes": path.stat().st_size,
            }
        )

    # ── Private parsers ────────────────────────────────────────

    def _parse_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {page_num}]\n{text.strip()}")
        if not pages:
            raise ValueError(f"No extractable text found in PDF: {path.name}. "
                             "File may be scanned/image-based.")
        return "\n\n".join(pages)

    def _parse_docx(self, path: Path) -> str:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError(f"No text content found in DOCX: {path.name}")
        return "\n\n".join(paragraphs)

    def _parse_txt(self, path: Path) -> str:
        content = path.read_text(encoding="utf-8", errors="replace").strip()
        if not content:
            raise ValueError(f"TXT file is empty: {path.name}")
        return content

    def _parse_csv(self, path: Path) -> str:
        df = pd.read_csv(path)
        if df.empty:
            raise ValueError(f"CSV file is empty: {path.name}")
        # Convert each row to a readable sentence-like string
        # This makes CSV content semantically searchable
        rows = []
        for _, row in df.iterrows():
            row_text = ", ".join(
                f"{col}: {val}"
                for col, val in row.items()
                if pd.notna(val)
            )
            rows.append(row_text)
        return "\n".join(rows)